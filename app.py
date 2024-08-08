from selenium import webdriver
from selenium.webdriver.common.by import By
from datetime import datetime
import pytz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx2pdf import convert

def obter_data_hora():
    """Obtém a data e hora no fuso horário de São Paulo."""
    try:
        saopaulo_tz = pytz.timezone('America/Sao_Paulo')
        saopaulo_time = datetime.now(saopaulo_tz)
        return saopaulo_time.strftime('%d/%m/%Y')
    except Exception as e:
        print(f"Erro ao obter a data e hora: {e}")
        return 'Data/Hora não disponível'

def fazer_print(driver):
    """Faz um print da tela e salva a imagem."""
    print("Fazendo um print da tela...")
    screenshot_path = 'screenshot.png'
    driver.save_screenshot(screenshot_path)
    print(f"Print da tela salvo em {screenshot_path}")

def extrair_informacoes(driver):
    """Extrai informações da página e retorna data, cotação e URL."""
    print("Extraindo informações da página...")
    element_cotacao_real = driver.find_elements(By.XPATH, '//*[@id="root"]/div[2]/div/div[1]/div/div[3]/div[2]/form/div[2]/input')
    data_hora = obter_data_hora()
    url_cotacao = driver.current_url

    if element_cotacao_real:
        cotacao_real = element_cotacao_real[0].get_attribute('value')
        return data_hora, cotacao_real, url_cotacao[8:]
    else:
        print('Nenhum elemento input encontrado')
        return data_hora, 'N/A', 'N/A'

def criar_word(titulo, cotacao, data, url, caminho_imagem, autor):
    """Cria um documento Word com as informações fornecidas."""
    print("Criando documento Word...")
    doc = Document()

    titulo_word = doc.add_heading(f'{titulo} - {cotacao} ({data})', level=1)
    titulo_word.alignment = 1
    run = titulo_word.runs[0]
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.underline = None

    parrafo_1 = doc.add_paragraph(f"O dólar está no valor de {cotacao}, na data {data}.")
    parrafo_1.paragraph_format.space_after = Pt(12) 

    parrafo_2 = doc.add_paragraph()
    run = parrafo_2.add_run("Valor cotado no site ")
    run.font.color.rgb = RGBColor(0, 0, 0)
    link_run = parrafo_2.add_run(url)
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    link_run.font.underline = True
    parrafo_2.paragraph_format.space_after = Pt(12)

    parrafo_3 = doc.add_paragraph("Print da cotação atual:")
    parrafo_3.paragraph_format.space_after = Pt(12)
    doc.add_picture(caminho_imagem, width=Inches(6), height=Inches(3))

    texto_4 = f"Cotação feita por – {autor}."
    doc.add_paragraph(texto_4)

    doc.save('cotacao_dolar.docx')
    print(f"Documento Word salvo em cotacao_dolar.docx")

def converter_word_para_pdf(arquivo_docx, arquivo_pdf):
    """Converte o documento Word para PDF."""
    print(f"Convertendo {arquivo_docx} para PDF...")
    convert(arquivo_docx, arquivo_pdf)
    print(f'Arquivo PDF gerado: {arquivo_pdf}')

def main():
    """Função principal que executa o fluxo do programa."""
    print("Iniciando o processo...")

    options = webdriver.ChromeOptions()
    options.add_argument("--no-sandbox")
    options.add_argument("--disable-dev-shm-usage")
    options.add_argument("--disable-notifications")
    options.add_argument("--disable-popup-blocking")
    options.add_argument("--start-maximized")
    options.add_argument("--disable-infobars")

    driver = webdriver.Chrome(options=options)

    url = 'https://www.remessaonline.com.br/cotacao/cotacao-dolar'
    driver.get(url)
    driver.implicitly_wait(10)
    print(f"Acessando o site: {url}")

    autor = 'Angel Cruz'
    titulo = 'Cotação Atual do Dólar'
    caminho_imagem = 'screenshot.png'
    caminho_word = 'cotacao_dolar.docx'

    # Extrair informações
    data_hora, cotacao_real, url_cotacao = extrair_informacoes(driver)

    # Fazer print da tela
    fazer_print(driver)

    # Criar o documento Word
    criar_word(titulo, cotacao_real, data_hora, url_cotacao, caminho_imagem, autor)

    # Converter para PDF
    converter_word_para_pdf(caminho_word, 'cotacao_dolar.pdf')

    # Fechar o navegador
    driver.quit()
    print("Processo concluído.")

if __name__ == '__main__':
    main()
